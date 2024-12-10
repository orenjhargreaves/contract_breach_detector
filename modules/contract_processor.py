from docx import Document
from typing import List, Tuple

from .query_llm import QueryLLM
from .evidence_search import EvidenceSearch


class ContractProcessor:
    """
    A class for processing contract documents to extract terms and generate HTML highlights
    using an LLM (Large Language Model).

    Attributes:
        llm (QueryLLM): An instance of the LLM interface for querying and extracting data.
    """

    def __init__(self, llm: QueryLLM, evidence_search: EvidenceSearch):
        """
        Initializes the ContractProcessor class.

        Args:
            llm (QueryLLM): The LLM instance for querying and data extraction.
        """
        self.llm = llm
        self.evidence_search = evidence_search

    def load_document(self, filepath: str) -> Document:
        """
        Loads a .docx document from the specified file path.

        Args:
            filepath (str): The path to the document file.

        Returns:
            Document: A `Document` object representing the loaded .docx file.
        """
        return Document(filepath)

    def extract_terms(self, document: Document, terms: dict) -> dict:
        """
        Extracts contract terms from a document using the LLM.

        Args:
            document (Document): A `Document` object representing the contract.
            terms (dict): A dictionary specifying the format of terms to extract.

        Returns:
            dict: The extracted terms in JSON format.
        """
        # Combine all paragraphs into a single string
        text = "\n".join([p.text for p in document.paragraphs])

        messages = [
            {
                "role": "system",
                "content": (
                    f"You are an AI assistant tasked with extracting structured information from documents. "
                    f"Output the extracted information in the following JSON format: {terms}."
                ),
            },
            {
                "role": "user",
                "content": f"Extract the key details from the following document:\n\n{text}",
            },
        ]

        # Query the LLM
        response = self.llm.query_llm(messages)
        return response

    def extract_terms_with_locations(self, document: Document, fields: list) -> dict:
        """
        Extracts specific fields along with their locations in the document.

        Args:
            document (Document): A `Document` object representing the contract.
            fields (list): A list of fields to extract with their start and end positions.

        Returns:
            dict: The extracted fields with their values and locations in JSON format.
        """
        # Combine all paragraphs into a single string
        text = "\n".join([p.text for p in document.paragraphs])

        messages = [
            {
                "role": "system",
                "content": (
                    f"You are an AI assistant document parser. Extract the specified fields along with their "
                    f"locations in the document. Each field should have a 'value', 'start_position', and 'end_position'. "
                    f"The fields are: {fields}. Output in the JSON format: "
                    f"{{'field': {{'value': '', 'start_position': '', 'end_position': ''}}...}}."
                ),
            },
            {
                "role": "user",
                "content": (
                    f"Extract the key details from the following document. If a value is missing, leave it empty. "
                    f"Ensure the response follows the described format:\n\n{text}"
                ),
            },
        ]

        # Query the LLM
        response = self.llm.query_llm(messages)
        return response

    def generate_html_highlight(self, document: Document, annotations: dict, output_path: str):
        """
        Generates an HTML file with highlighted terms from the document.

        Args:
            document (Document): A `Document` object representing the contract.
            annotations (dict): A dictionary of extracted fields with start and end positions.
            output_path (str): The file path to save the generated HTML file.
        """
        # Combine all paragraphs into a single string
        text = "\n".join([p.text for p in document.paragraphs])

        # Sort extracted fields by their start positions
        sorted_fields = sorted(
            annotations.values(),
            key=lambda x: int(x['start_position']) if x['start_position'] else float('inf'),
        )

        # Highlight extracted fields in the text
        offset = 0
        for field in sorted_fields:
            if field['value']:  # Only highlight fields with valid values
                start = int(field['start_position']) + offset
                end = int(field['end_position']) + offset
                text = (
                    text[:start]
                    + f'<span style="background-color: yellow;">{text[start:end]}</span>'
                    + text[end:]
                )
                offset += len('<span style="background-color: yellow;"></span>')

        # Save the highlighted text as an HTML file
        with open(output_path, "w") as file:
            file.write(f"<html><body><pre>{text}</pre></body></html>")

    def generate_html_highlight_from_fuzz(self, document: Document, list_of_anno: List[Tuple[int,int]], output_path: str):
        """
        Generates an HTML file with highlighted terms from the document.

        Args:
            document (Document): A `Document` object representing the contract.
            annotations (dict): A dictionary of extracted fields with start and end positions.
            output_path (str): The file path to save the generated HTML file.
        """
        # Combine all paragraphs into a single string
        text = "\n".join([p.text for p in document.paragraphs])
        
        # Sort annotations by start position to ensure proper order
        list_of_anno = sorted(list_of_anno, key=lambda x: x[0])

        # Highlight extracted fields in the text
        offset = 0  # Tracks the additional length added by the HTML tags
        for start, end in list_of_anno:
            if start == -1 & end == -1:
                continue
            start += offset
            end += offset
            highlight = f'<span style="background-color: yellow;">{text[start:end]}</span>'
            text = text[:start] + highlight + text[end:]
            offset += len(highlight) - (end - start)

        # Save the highlighted text as an HTML file
        with open(output_path, "w") as file:
            file.write(f"<html><body><pre>{text}</pre></body></html>")

    # def check_extracted_terms(self, )