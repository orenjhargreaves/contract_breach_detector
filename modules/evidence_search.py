
from .query_llm import QueryLLM
from docx import Document
from rapidfuzz import process, fuzz, distance
from difflib import SequenceMatcher
from fuzzysearch import find_near_matches
from typing import Tuple

class EvidenceSearch:
    """
    Will search for evidence within a document
    """

    def __init__(self, llm: QueryLLM):
        """initialises Evidence Search class
        
        Args:
            llm (QueryLLM): An instance of the LLM interface for analysis.
            """
        
        self.llm = llm
    
    def find_exact_text(self, text: str, document: Document):
        """
        when passed text that is believed to have been evidenced from a document, it will return the exact text form within the document.
        Args:
            text (str): the text to find in document
            document (Document): the document within to search
        """
        document_text = "\n".join([p.text for p in document.paragraphs])

        messages = [
            {
                "role": "system",
                "content": (
                    "You are an AI assistent that when passed in a text string you will search the given"
                    "document for the relavent information and return the exact wording from the document."
                    "in the documentThe input from the user will come in the form:"
                    "Document <document text>"
                    "text to evidence: <text>"
                    "Output in the JSON form {'evidence': <exact wording>}"
                    ""
                    "Document: 'This is a test contract. specific piece of information to be returned: password123. Rest of the test contract'"
                    "text to evidence: 'specific piece of information: password123'"
                    "{'evidence': 'specific piece of information to be returned: password123'"
                ),
            },
            {
                "role": "user",
                "content": (
                    f"Document: {document_text}"
                    f"text to evidence: {text}"
                    "{'evidence': "
                ),
            },
        ]

        # Query the LLM
        response = self.llm.query_llm(messages)
        return response
    
    def find_best_fuzzy_match(self, text:str, document: Document, threshold=80):
        """
        when given text and a document it will find the location of the best match
        Args:
            text (str): the text to find in the document
            document (Document): the document in which to find text
        """
        document_text = "\n".join([p.text for p in document.paragraphs])
        # Use rapidfuzz to calculate similarity score
        score = fuzz.token_set_ratio(text, document_text)
        
        # Check if score meets the threshold
        if score >= threshold:
            # Locate the start and end indexes of the best match
            start_index = document_text.find(text)
            if start_index != -1:
                end_index = start_index + len(text)
            else:
                end_index = -1
            return score, start_index, end_index
        else:
            return None
        
    def find_best_match_in_document(self, text: str, document: Document, threshold=50)-> Tuple[str, int, int, int]:
        """
        splitting find_best_fuzzy_match to just find the match here.
        Args:
            text (str): the text to find in the document
            document (Document): the document in which to find text
        Returns:
            (matched_text, similarity_score, start_index, end_index) where start_index and end_index are the positions of matched_text in the document
        """
        document_text = "\n".join([p.text for p in document.paragraphs])
        # print(document_text)
        best_match = process.extractOne(text, document_text, scorer=fuzz.token_set_ratio, score_cutoff=threshold)

        if best_match:
            matched_text, similarity_score, _ = best_match
            start_index = document_text.find(matched_text)
            end_index = start_index + len(matched_text)
            return (matched_text, similarity_score, start_index, end_index)
        else:
            return (None, None,None, None)
    
    def calculate_levenshtein_distance(text:str, document_text:str)->float:
        """
        Takes two matching pieces of text and calculates the Levenshtein distance between them.
        Args:
            text (str): the text to find in the document
            document_text (str): the best matched text in the document
        """
        levenshtein_distance = distance.Levenshtein.distance(text, document_text)
        ld_score = (max(len(text), len(document_text)) - levenshtein_distance)/max(len(text), len(document_text))
        return ld_score
    
    def find_best_match_in_document_v2(self, text: str, document: Document, threshold=50)-> Tuple[str, float, int, int]:
        """
        splitting find_best_fuzzy_match to just find the match here.
        Args:
            text (str): the text to find in the document
            document (Document): the document in which to find text
        Returns:
            (matched_text, similarity_score, start_index, end_index) where start_index and end_index are the positions of matched_text in the document
        """
        document_text = "\n".join([p.text for p in document.paragraphs])
        # print(document_text)
        # best_match = process.extractOne(text, document_text, scorer=fuzz.token_set_ratio, score_cutoff=threshold)
        best_match = find_near_matches(text, document_text, max_l_dist=max(2, int(len(text)*0.5)))
        if len(best_match)!= 0:
            similarity_score = (len(text) - best_match[0].dist)/len(text)
            return (best_match[0].matched, similarity_score, best_match[0].start, best_match[0].end)
        else:
            return (None, None, None, None)

        # if best_match:
        #     matched_text, similarity_score, _ = best_match
        #     start_index = document_text.find(matched_text)
        #     end_index = start_index + len(matched_text)
        #     return (matched_text, similarity_score, start_index, end_index)
        # else:
        #     return (None, None,None, None)