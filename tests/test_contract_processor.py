import sys
import os
import unittest
from unittest.mock import MagicMock
from docx import Document
from contract_breach_detector.modules.contract_processor import ContractProcessor
from contract_breach_detector.modules.query_llm import QueryLLM


class TestContractProcessor(unittest.TestCase):
    """
    Test suite for the ContractProcessor class.
    """

    def setUp(self):
        """
        Sets up the test environment by mocking the QueryLLM instance
        and creating a ContractProcessor instance.
        """
        self.mock_llm = MagicMock(spec=QueryLLM)
        self.processor = ContractProcessor(self.mock_llm)
        self.test_document_path = os.path.join(os.path.dirname(__file__), "test_data", "test_contract.docx")
        self.test_html_output_path = os.path.join(os.path.dirname(__file__), "test_data", "highlighted_contract.html")
        self.test_annotations = {
            "field1": {"value": "test", "start_position": "10", "end_position": "14"},
            "field2": {"value": "highlighted", "start_position": "29", "end_position": "40"},
        }

    def test_load_document(self):
        """
        Tests the load_document method to ensure it correctly loads a .docx file.
        """
        # Create a temporary test document
        doc = Document()
        doc.add_paragraph("This is a test contract.")
        doc.save(self.test_document_path)

        loaded_doc = self.processor.load_document(self.test_document_path)
        self.assertIsInstance(loaded_doc, type(Document()))
        self.assertEqual(loaded_doc.paragraphs[0].text, "This is a test contract.")

        # Clean up the temporary file
        os.remove(self.test_document_path)

    def test_extract_terms(self):
        """
        Tests the extract_terms method to ensure it correctly queries the LLM and returns the expected response.
        """
        # Mock LLM response
        mock_response = {"contract_number": "12345", "delivery_date": "2024-06-15"}
        self.mock_llm.query_llm.return_value = mock_response

        # Create a test document
        doc = Document()
        doc.add_paragraph("Contract Number: 12345\nDelivery Date: 2024-06-15")

        terms_format = {"contract_number": "", "delivery_date": ""}
        result = self.processor.extract_terms(doc, terms_format)

        self.mock_llm.query_llm.assert_called_once()  # Ensure LLM was called
        self.assertEqual(result, mock_response)  # Ensure response matches mock

    def test_extract_terms_with_locations(self):
        """
        Tests the extract_terms_with_locations method to ensure it correctly queries the LLM and returns field locations.
        """
        # Mock LLM response
        mock_response = {
            "contract_number": {"value": "12345", "start_position": "10", "end_position": "15"},
            "delivery_date": {"value": "2024-06-15", "start_position": "25", "end_position": "35"},
        }
        self.mock_llm.query_llm.return_value = mock_response

        # Create a test document
        doc = Document()
        doc.add_paragraph("Contract Number: 12345\nDelivery Date: 2024-06-15")

        fields = ["contract_number", "delivery_date"]
        result = self.processor.extract_terms_with_locations(doc, fields)

        self.mock_llm.query_llm.assert_called_once()  # Ensure LLM was called
        self.assertEqual(result, mock_response)  # Ensure response matches mock

    def test_generate_html_highlight(self):
        """
        Tests the generate_html_highlight method to ensure it correctly generates an HTML file with highlights.
        """
        # Create a test document
        doc = Document()
        doc.add_paragraph("This is a test contract with highlighted fields.")

        # Run the method
        self.processor.generate_html_highlight(doc, self.test_annotations, self.test_html_output_path)

        # Verify the HTML file is created
        self.assertTrue(os.path.exists(self.test_html_output_path))

        # Verify the content contains the highlights
        with open(self.test_html_output_path, "r") as file:
            content = file.read()
            self.assertIn('<span style="background-color: yellow;">test</span>', content)
            self.assertIn('<span style="background-color: yellow;">highlighted</span>', content)

        # Clean up the test HTML file
        os.remove(self.test_html_output_path)


if __name__ == "__main__":
    unittest.main()
