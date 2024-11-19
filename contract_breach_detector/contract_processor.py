from docx import Document
import openai
import os
import pickle
import json
import hashlib

# extracted_terms_format = {
#   "contract_details": {
#     "contract_number": "",
#     "commencement_date": "",
#     "parties": {
#       "customer": {
#         "name": "",
#         "address": "",
#         "representative": ""
#       },
#       "supplier": {
#         "name": "",
#         "address": "",
#         "vat_number": "",
#         "representative": ""
#       }
#     }
#   },
#   "goods_details": {
#     "description": "",
#     "quantity": "",
#     "specifications": {
#       "dimensions": "",
#       "material_composition": "",
#       "quality_standards": "",
#       "surface_finish": "",
#       "mechanical_properties": "",
#       "flatness_tolerance": "",
#       "compliance_and_certifications": {
#         "regulatory_compliance": "",
#         "certification": ""
#       }
#     }
#   },
#   "delivery": {
#     "delivery_date": "",
#     "delivery_location": "",
#     "packaging": {
#       "palletization": "",
#       "labeling": "",
#       "handling_instructions": ""
#     },
#     "penalties": {
#       "late_delivery": {
#         "penalty_rate": "",
#         "maximum_penalty": "",
#         "termination_threshold": ""
#       }
#     }
#   },
#   "price_and_payment": {
#     "contract_price": "",
#     "payment_terms": "",
#     "taxes_and_duties": "",
#     "set_off_rights": ""
#   },
#   "warranties": {
#     "warranty_period": "",
#     "warranty_conditions": "",
#     "remedies_for_breach": ""
#   },
#   "liability_and_indemnification": {
#     "limitation_of_liability": "",
#     "cap_on_liability": "",
#     "supplier_indemnification": "",
#     "insurance": ""
#   },
#   "termination": {
#     "termination_for_convenience": "",
#     "termination_for_cause": {
#       "breach_period": "",
#       "bankruptcy": ""
#     },
#     "consequences_of_termination": ""
#   },
#   "force_majeure": {
#     "definition": "",
#     "effect": "",
#     "obligations": "",
#     "termination_due_to_prolonged_event": ""
#   },
#   "confidentiality": {
#     "scope": "",
#     "obligations": "",
#     "exclusions": "",
#     "return_or_destruction": ""
#   },
#   "intellectual_property": {
#     "ownership": "",
#     "license_grant": "",
#     "indemnity": ""
#   },
#   "dispute_resolution": {
#     "negotiation": "",
#     "mediation": "",
#     "arbitration": {
#       "rules": "",
#       "seat_of_arbitration": "",
#       "language": ""
#     }
#   },
#   "governing_law": "",
#   "notices": {
#     "customer_address": "",
#     "supplier_address": "",
#     "effective_date": ""
#   },
#   "schedules": {
#     "schedule_1_goods_specifications": "",
#     "schedule_2_mandatory_policies": ""
#   }
# }


class ContractProcessor:
    def __init__(self, model: str = "gpt-4o-mini", cache_dir: str = "/home/oren/Documents/Python/magentic_poc/contract_breach_detector/cache"):
        self.model = model
        self.client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        self.cache_dir = cache_dir
        os.makedirs(self.cache_dir, exist_ok=True)  # Ensure the cache directory exists

    def load_document(self, filepath: str) -> Document:
        """Loads document form given filepath. uses docx type Document"""
        return Document(filepath)
    
    def _generate_hash(self, raw_text: str) -> str:
        """
        Generate a unique hash for the document text and prompt.
        """
        return hashlib.sha256(raw_text.encode()).hexdigest()

    def _cache_path(self, hash_key: str) -> str:
        """
        Get the file path for the cached response.
        """
        return os.path.join(self.cache_dir, f"{hash_key}.pkl")

    def _load_from_cache(self, hash_key: str):
        """
        Load a response from the cache if it exists.
        """
        cache_file = self._cache_path(hash_key)
        if os.path.exists(cache_file):
            with open(cache_file, "rb") as file:
                return pickle.load(file)
        return None

    def _save_to_cache(self, hash_key: str, response):
        """
        Save a response to the cache.
        """
        cache_file = self._cache_path(hash_key)
        with open(cache_file, "wb") as file:
            pickle.dump(response, file)

    def extract_terms(self, document: Document, terms: json) -> dict:
        """
        Extract contract terms from a document using LLM or cache.
        """
        text = "\n".join([p.text for p in document.paragraphs])
        
        messages = [
                    {"role": "system", "content": f"You are an AI assistant that extracts structured information from documents and outputs it in the JSON format: {terms}"},
                    {"role": "user", "content": f"Extract the key details from the following document and format them as a JSON object:\n\n{text}"}
        ]
        # Generate a hash for the message
        query_hash = self._generate_hash(f"{terms}, {text}")

        # Check if the response is cached
        cached_response = self._load_from_cache(query_hash)
        if cached_response:
            print("Loaded response from cache.")
            return cached_response

        response = self.client.chat.completions.create(model=self.model, messages=messages)
        
        # Parse the response
        extracted_terms = json.loads(response.choices[0].message.content[7:-4])

        # Cache the response
        self._save_to_cache(query_hash, extracted_terms)
        
        return extracted_terms
    
    def extract_terms_with_locations(self, document: Document, fields: list) -> dict:
        text = "\n".join([p.text for p in document.paragraphs]) 
        messages = [
                    {"role": "system", "content": f"You are a document parser extracting key information from contracts. Analyze the following document and extract the specified fields along with the locations where they were found the 'start_position' must be at the start of the information sumarised by the 'value' and the 'end_position' after. The fields are: {fields}. For each 'field' in fields, return a dictionary of the result in this format: {{'field'{{'value': 'value', 'start_position': start, 'end_position': end}}...}}"},
                    {"role": "user", "content": f"Extract the key details from the following document and format them as described above. do not make up any values. if you are unsure leave the, blank. Any fields you fill, you must reference their start and end locations:\n\n{text}"}
        ]

        key = " ".join(f"{message['role']}: {message['content']}" for message in messages)

        # Generate a hash for the message
        query_hash = self._generate_hash(key)

        # Check if the response is cached
        cached_response = self._load_from_cache(query_hash)
        if cached_response:
            print("Loaded response from cache.")
            return cached_response

        response = self.client.chat.completions.create(model=self.model, messages=messages)
        
        # Parse the response
        extracted_terms = json.loads(response.choices[0].message.content[7:-4])

        # Cache the response
        self._save_to_cache(query_hash, extracted_terms)
        
        return extracted_terms
    
    def generate_html_highlight(self, document: Document, annotations: dict, output_path: str):
        # Generate text from document
        text = "\n".join([p.text for p in document.paragraphs])
        # Sort extracted fields by start_position
        sorted_fields = sorted(annotations.values(), key=lambda x: int(x['start_position']) if x['start_position'] else float('inf'))

        # Adjust text with HTML tags for highlights
        offset = 0
        for field in sorted_fields:
            if field['value']:  # Skip empty fields
                start = field['start_position'] + offset
                end = field['end_position'] + offset
                text = (
                    text[:start]
                    + f'<span style="background-color: yellow;">{text[start:end]}</span>'
                    + text[end:]
                )
                offset += len('<span style="background-color: yellow;"></span>')

        # Save as HTML
        with open(output_path, "w") as file:
            file.write(f"<html><body><pre>{text}</pre></body></html>")
